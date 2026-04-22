from typing import List 
import pdfplumber 
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

            
class ResumeParser:
    def __init__(self, file_path:str):
        self.file_path = file_path

    def parse_pdf(self) ->str:
        """Logic for coordinate-based PDF density parsing."""

        body_text_segments:List[str] = []

        with pdfplumber.open(self.file_path) as pdf:
            for page in pdf.pages:
                # --- DENSITY LOGIC START ---
                num_slices = 20 
                slice_width = page.width / num_slices
                density_map = []

                for i in range(num_slices):
                    bbox = (i * slice_width, 0, (i + 1) * slice_width, page.height)
                    words = page.crop(bbox).extract_words()
                    density_map.append(sum(len(w["text"]) for w in words))

                mid_start, mid_end = 4, 16
                min_density = min(density_map[mid_start:mid_end])
                min_indices = [i for i, d in enumerate(density_map) if d == min_density and mid_start <= i < mid_end]

                valley_start, valley_end = min_indices[0], min_indices[-1] + 1
                split_x = ((valley_start + valley_end) / 2) * slice_width

                active_slices = [d for d in density_map if d > 0]
                # --- DECISION LOGIC ---
                if active_slices:
                    avg_active = sum(active_slices) / len(active_slices)
                    if min_density < (avg_active * 0.15):
                        left = page.crop((0, 0, split_x, page.height)).extract_text() or ""
                        right = page.crop((split_x, 0, page.width, page.height)).extract_text() or ""
                        body_text_segments.append(f"{left}\n{right}")
                        continue # Move to next page
                
                # Fallback for single column
                body_text_segments.append(page.extract_text() or "")
        
        return "\n".join(body_text_segments)

    def parse_docx(self) -> str:
        """Logic for structural XML parsing of Word documents."""
        doc = Document(self.file_path)
        combined_data:List[str] = []

        # 1. Main Flow (Paragraphs and Tables)
        # Note: We need to use internal element checks
        for element in doc.element.body:
            # Reconstruct objects from XML elements
            if element.tag.endswith('p'):
                para = Paragraph(element, doc)
                if para.text.strip():
                    combined_data.append(para.text.strip())
            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        combined_data.append(row_text)

        # 2. Text Box Hunter (The 'Spam' protection)
        seen_texts = set()
        box_texts = []
        for txbx in doc.element.xpath('.//*[local-name()="txbxContent"]'):
            current_box = "\n".join(Paragraph(p, doc).text.strip() for p in txbx.xpath('.//*[local-name()="p"]') if Paragraph(p, doc).text.strip())
            if current_box and current_box not in seen_texts:
                box_texts.append(current_box)
                seen_texts.add(current_box)

        if box_texts:
            combined_data.append("\n--- ADDITIONAL DATA ---")
            combined_data.extend(box_texts)

        return "\n".join(combined_data)

    def resume_text_extractor(self) ->str:
        """The 'Brain' that decides which parser to use."""
        if self.file_path.lower().endswith(".pdf"):
            return self.parse_pdf()
        elif self.file_path.lower().endswith(".docx"):
            return self.parse_docx()
        else:
            return "" # add a feature to add name of file which is failed in parsing 


    # header parser 
    def get_pdf_header_slice(self, slice_percentage=0.20):

        header_text = ""
        
        try:
            with pdfplumber.open(self.file_path) as pdf:
                if not pdf.pages:
                    return ""
                
                first_page = pdf.pages[0]
                
                # Define the bounding box: (x0, top, x1, bottom)
                # pdfplumber uses points (1/72 inch) starting from top-left (0,0)
                width = first_page.width
                height = first_page.height
                
                bbox = (0, 0, width, height * slice_percentage)
                
                # Crop the page to the header area
                header_area = first_page.within_bbox(bbox)
                
                # Extract text from the cropped area
                header_text = header_area.extract_text()
                
        except Exception as e:
            print(f"Extraction Error: {e}")
            return ""
            
        return header_text if header_text else ""

    def extract_docx_header(self,line_limit=10):

        combined_header_parts = []
        
        try:
            doc = Document(self.file_path)
            
            # 1. Extract from the Formal XML Header parts
            # Documents can have multiple sections, each with its own header
            for section in doc.sections:
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            combined_header_parts.append(text)
            
            # 2. Extract from the top of the Document Body
            # In case the user didn't use the formal 'Header' feature
            body_paragraphs = doc.paragraphs[:line_limit]
            for para in body_paragraphs:
                text = para.text.strip()
                if text:
                    combined_header_parts.append(text)
                    
        except Exception as e:
            # Log error for debugging, but return empty string to keep pipeline moving
            print(f"DOCX Extraction Error: {e}")
            return ""
        
        # Join with newlines and return
        return "\n".join(combined_header_parts).strip()

    def main_header_extractor(self):
        if self.file_path.lower().endswith(".pdf"):
            return self.get_pdf_header_slice()
        elif self.file_path.lower().endswith(".docx"):
            return self.extract_docx_header()
        else:
            return ""